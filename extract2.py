import docx
import fitz
import sys

def extract_docx_with_tables(filepath, outpath):
    try:
        doc = docx.Document(filepath)
        text = []
        for element in doc.element.body:
            if element.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._element == element:
                        text.append(para.text)
                        break
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                row_data.append(cell.text.replace("\n", " ").strip())
                            text.append(" | ".join(row_data))
                        break
        with open(outpath, "w", encoding="utf-8") as f:
            f.write("\n".join(text))
        print(f"Extracted DOCX to {outpath}")
    except Exception as e:
        print(f"Error extracting docx: {e}")

def convert_pdf_to_images(filepath):
    try:
        doc = fitz.open(filepath)
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=150)
            pix.save(f"prd_page_{i}.png")
        print(f"Converted PDF to {len(doc)} images.")
    except Exception as e:
        print(f"Error converting pdf to images: {e}")

if __name__ == "__main__":
    extract_docx_with_tables("CampusTrade_TRD.docx", "CampusTrade_full.txt")
    convert_pdf_to_images("prd.pdf")
